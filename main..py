import glob
import time
import telebot
from docx import Document
from telebot.apihelper import ApiTelegramException
import io
import Levenshtein
import sqlite3
import os
import requests 

bot = telebot.TeleBot('6303559793:AAGysyRYVE0v_WCROLLDYXh9ApFbm4CmhYo')
ADMIN_CHAT_ID = 2099795903

# Директория для хранения файла расписания
SCHEDULE_DIR = 'schedule_files'
if not os.path.exists(SCHEDULE_DIR):
    os.makedirs(SCHEDULE_DIR)

def save_schedule_file(downloaded_file):
    schedule_file_path = os.path.join(SCHEDULE_DIR)
    
    # Удаление старого файла расписания
    if os.path.exists(schedule_file_path):
        os.remove(schedule_file_path)
    
    # Сохранение нового файла расписания
    with open(schedule_file_path, 'wb') as f:
        f.write(downloaded_file.read())

def load_latest_schedule():
    list_of_files = glob.glob(os.path.join(SCHEDULE_DIR, '*.docx')) 
    latest_file = max(list_of_files, key=os.path.getctime)
    print(f"Loading schedule from {latest_file}")
    with open(latest_file, 'rb') as f:
        return f.read()

def load_schedule_file():
    schedule_file_path = os.path.join(SCHEDULE_DIR)
    
    if os.path.exists(schedule_file_path):
        print(f"Loading schedule from {schedule_file_path}")  # Добавлено для диагностики
        with open(schedule_file_path, 'rb') as f:
            return f.read()
    else:
        print(f"Schedule file {schedule_file_path} not found")  # Добавлено для диагностики
    return None


schedule_dict = {}
teacher_schedule_dict = {}
feedback_users = {}
broadcast_mode = False
# Инициализация базы данных
def init_db():
    conn = sqlite3.connect('users.db')
    cursor = conn.cursor()
    cursor.execute('''CREATE TABLE IF NOT EXISTS users
                      (chat_id INTEGER PRIMARY KEY, first_name TEXT, last_name TEXT, username TEXT)''')
    conn.commit()
    conn.close()

init_db()


def init_specialties_db():
    conn = sqlite3.connect('specialties.db')
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS specialties (
            id INTEGER PRIMARY KEY,
            name TEXT,
            qualification TEXT,
            study_duration TEXT,
            study_form TEXT,
            study_language TEXT,
            description TEXT
        )
    ''')
    conn.commit()
    conn.close()

init_specialties_db()

def add_or_update_user(chat_id, first_name, last_name, username):
    conn = sqlite3.connect('users.db')
    cursor = conn.cursor()
    cursor.execute("INSERT OR REPLACE INTO users (chat_id, first_name, last_name, username) VALUES (?, ?, ?, ?)", (chat_id, first_name, last_name, username))
    conn.commit()
    conn.close()

def safe_send_message(chat_id, message_text):
    try:
        bot.send_message(chat_id, message_text)
    except ApiTelegramException as e:
        if e.result.status_code == 403:
            print(f'User {chat_id} has blocked the bot.')

def safe_reply_to(message, text):
    try:
        bot.reply_to(message, text)
    except ApiTelegramException as e:
        if e.result.status_code == 403:
            print(f'User {message.chat.id} has blocked the bot.')

def insert_specialty(name, qualification, study_duration, study_form, study_language, description):
    conn = sqlite3.connect('specialties.db')
    cursor = conn.cursor()
    cursor.execute("INSERT INTO specialties (name, qualification, study_duration, study_form, study_language, description) VALUES (?, ?, ?, ?, ?, ?)",
                   (name, qualification, study_duration, study_form, study_language, description))
    conn.commit()
    conn.close()

def get_specialty_by_name(name):
    conn = sqlite3.connect('specialties.db')
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM specialties WHERE name=?", (name,))
    data = cursor.fetchone()
    conn.close()
    return data

def insert_specialty_data(specialty_data):
    """
    Inserts specialty data into the database.
    """
    conn = sqlite3.connect('specialties.db')
    cursor = conn.cursor()
    cursor.execute("INSERT INTO specialties (name, qualification, study_duration, study_form, study_language, description) VALUES (?, ?, ?, ?, ?, ?)",
                   (specialty_data["name"], specialty_data["qualification"], specialty_data["study_duration"], specialty_data["study_form"], specialty_data["study_language"], specialty_data["description"]))
    conn.commit()
    conn.close()

def fetch_all_specialties():
    """
    Fetches all specialties from the database.
    """
    conn = sqlite3.connect('specialties.db')
    cursor = conn.cursor()
    cursor.execute("SELECT name FROM specialties")
    specialties = cursor.fetchall()
    conn.close()
    return [item[0] for item in specialties]

# Sample data for testing




# Insert sample data for testing

# Fetch all specialties for testing
fetch_all_specialties()



def improved_parse_schedule(docx_file):
    doc = Document(docx_file)
    time_intervals = [cell.text.strip() for cell in doc.tables[0].rows[0].cells][1::2]
    schedule = {}
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            room = cells.pop(0)

            # Process data in pairs (Teacher, Group)
            for i in range(0, len(cells) - 1, 2):
                teachers = cells[i].split("\n") if i < len(cells) else []
                groups = cells[i+1].split("\n") if (i+1) < len(cells) else []
                
                for teacher, group in zip(teachers, groups):
                    if not teacher or not group:
                        continue

                    if group not in schedule:
                        schedule[group] = {}
                    
                    interval_index = i // 2
                    current_interval = time_intervals[interval_index]

                    schedule[group][current_interval] = {'room': room, 'teacher': teacher}

                    # Populate teacher_schedule_dict
                    if teacher not in teacher_schedule_dict:
                        teacher_schedule_dict[teacher] = {}
                    
                    teacher_schedule_dict[teacher][current_interval] = {'room': room, 'group': group}

    return schedule


loaded_schedule = load_latest_schedule()
if loaded_schedule:
    with io.BytesIO(loaded_schedule) as docx_file:
        schedule_dict = improved_parse_schedule(docx_file)

def find_closest_key(input_str, possible_keys):
    if not possible_keys:
        return None

    distances = [(key, Levenshtein.distance(input_str, key)) for key in possible_keys]
    closest_key = min(distances, key=lambda tup: tup[1])[0]
    return closest_key

@bot.message_handler(commands=['start', 'help'])
def send_welcome(message):
    add_or_update_user(message.chat.id, message.from_user.first_name, message.from_user.last_name, message.from_user.username)
    markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True)
    btn1 = telebot.types.KeyboardButton('Узнать расписание')
    btn2 = telebot.types.KeyboardButton('Оставить отзыв')
    btn3 = telebot.types.KeyboardButton('Информация про колледж') 
    markup.add(btn1, btn2, btn3)
    
    instruction = """
Привет! Это бот Политехнического колледжа. 

Вам доступны следующие функции:
1. Нажмите 'Узнать расписание' для получения расписания.
2. Нажмите 'Оставить отзыв', чтобы поделиться своим мнением или предложением.
3. Нажмите 'Информация про колледж', чтобы узнать информацию о специальностях.

Что вы хотите сделать?
"""
    safe_send_message(message.chat.id, instruction)

def notify_all_users(message_text):
    conn = sqlite3.connect('users.db')
    cursor = conn.cursor()
    for row in cursor.execute('SELECT chat_id FROM users'):
        chat_id = row[0]
        try:
           safe_send_message(chat_id, message_text)
        except ApiTelegramException as e:
            if e.result.status_code == 403:
                print(f"Пользователь {chat_id} заблокировал бота. Пропускаем.")
                continue  # Пропустить этого пользователя и продолжить цикл
            else:
                raise e  # Если это другая ошибка, поднимаем исключение
    conn.close()

@bot.message_handler(content_types=['document'])
def handle_docs(message):
    add_or_update_user(message.chat.id, message.from_user.first_name, message.from_user.last_name, message.from_user.username)
    
    # Если это админ, позволим загрузить расписание
    if message.chat.id == ADMIN_CHAT_ID:
        global schedule_dict
        if message.document.mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            file_info = bot.get_file(message.document.file_id)
            downloaded_file = bot.download_file(file_info.file_path)

            # Удаляем старый файл расписания, если он существует
            for filename in os.listdir(SCHEDULE_DIR):
                if filename.endswith('.docx'):
                    os.remove(os.path.join(SCHEDULE_DIR, filename))

            # Сохраняем новый файл расписания
            new_schedule_path = os.path.join(SCHEDULE_DIR, f"schedule_{message.document.file_name}")
            with open(new_schedule_path, 'wb') as new_file:
                new_file.write(downloaded_file)

            with open(new_schedule_path, 'rb') as docx_file:
                schedule_dict = improved_parse_schedule(docx_file)

            safe_reply_to(message, "Расписание успешно обновлено!")
            notify_all_users("Расписание было обновлено администратором! Запросите актуальное расписание.")
        else:
             safe_reply_to(message, "Пожалуйста, загрузите документ в формате .docx.")
    else:
        safe_send_message(message.chat.id, "Только админ может загрузить новое расписание.")
        

@bot.message_handler(func=lambda message: message.text == 'Узнать расписание')
def ask_for_schedule(message):
    if not schedule_dict:
         safe_send_message(message.chat.id, "Расписание еще не было загружено администратором. Пожалуйста, повторите попытку позже.")
    else:
         safe_send_message(message.chat.id, "Введите название вашей группы или имя преподавателя:")
    

@bot.message_handler(func=lambda message: message.text == 'Оставить отзыв')
def leave_feedback(message):
    safe_send_message(message.chat.id, "Пожалуйста, введите ваш отзыв или предложение. Вы можете завершить ввод, отправив команду /end.")
    feedback_users[message.chat.id] = ""
    
@bot.message_handler(func=lambda m: m.chat.id in feedback_users and m.text != '/end')
def collect_feedback(message):
    feedback_users[message.chat.id] += message.text + "\n"
    

@bot.message_handler(func=lambda m: m.text == '/end' and m.chat.id in feedback_users)
def end_feedback(message):
    feedback = feedback_users.pop(message.chat.id)
    
    user_info = message.from_user
    user_name = user_info.first_name or ""
    if user_info.last_name:
        user_name += " " + user_info.last_name
    user_username = "@" + user_info.username if user_info.username else ""
    
    # Формируем сообщение с информацией о пользователе и его отзывом
    feedback_message = f"Обратная связь от {user_name} {user_username} (ID: {message.chat.id}):\n{feedback}"
    
    bot.send_message(ADMIN_CHAT_ID, feedback_message)
    safe_send_message(message.chat.id, "Спасибо за ваш отзыв!")
 

@bot.message_handler(commands=['broadcast'])
def start_broadcast_mode(message):
    global broadcast_mode
    if message.chat.id == ADMIN_CHAT_ID:
        broadcast_mode = True
        bot.reply_to(message, "Вы активировали режим рассылки. Все ваши следующие сообщения будут отправлены всем пользователям. Чтобы завершить рассылку, отправьте команду /endbroadcast.")
        return

@bot.message_handler(commands=['endbroadcast'])
def stop_broadcast_mode(message):
    global broadcast_mode
    if message.chat.id == ADMIN_CHAT_ID and broadcast_mode:
        broadcast_mode = False
        bot.reply_to(message, "Вы завершили режим рассылки.")
        return


@bot.message_handler(func=lambda m: broadcast_mode and m.chat.id == ADMIN_CHAT_ID)
def send_broadcast(message):
    notify_all_users(message.text)
    

@bot.message_handler(func=lambda message: message.text == 'Информация про колледж')
def college_info(message):
    markup = telebot.types.InlineKeyboardMarkup()
    # Fetch all specialties from the database
    specialties = fetch_all_specialties()
    print(specialties)
    for specialty in specialties:
        # Проверим длину callback_data перед использованием
        if len(specialty.encode('utf-8')) <= 64:
            button = telebot.types.InlineKeyboardButton(text=specialty, callback_data=specialty)
            markup.add(button)
        else:
            print(f"Warning: Callback data '{specialty}' is too long and will be skipped.")
    bot.send_message(message.chat.id, "Выберите специальность:", reply_markup=markup)
 

@bot.callback_query_handler(func=lambda call: True)
def specialty_callback(call):
    specialty_name = call.data
    specialty_data = get_specialty_by_name(specialty_name)
    if specialty_data:
        response = f"""
Специальность: {specialty_data[1]}
Квалификация: {specialty_data[2]}
Срок обучения: {specialty_data[3]}
Форма обучения: {specialty_data[4]}
Язык обучения: {specialty_data[5]}
Описание: {specialty_data[6]}
"""
        try:
            bot.edit_message_text(chat_id=call.message.chat.id, message_id=call.message.message_id, text=response)
        except ApiTelegramException as e:
            if e.result.status_code == 403:
                print(f'User {call.message.chat.id} has blocked the bot.')
    else:
        safe_send_message(call.message.chat.id, "Извините, информация о данной специальности отсутствует.")
    
  

@bot.message_handler(func=lambda m: True)
def send_schedule(message):
    if message.chat.id == ADMIN_CHAT_ID and (message.text == "/broadcast" or message.text == "/endbroadcast"):
        return
    query = message.text
    response = ""

    # Try to find group first
    if query in schedule_dict:
        schedule_for_group = schedule_dict[query]
        response = "Расписание для группы {}:".format(query)
        for _, details in schedule_for_group.items():
            teacher_info = details["teacher"] if details["teacher"] else "Преподаватель: Не указан"
            response += "\nАудитория: {} - {}".format(details["room"], teacher_info)
    elif query in teacher_schedule_dict:
        schedule_for_teacher = teacher_schedule_dict[query]
        response = "Расписание для преподавателя {}:".format(query)
        for _, details in schedule_for_teacher.items():
            group_info = details["group"] if details["group"] else "Группа: Не указана"
            response += "\nАудитория: {} - {}".format(details["room"], group_info)
    else:
        closest_group = find_closest_key(query, schedule_dict.keys())
        closest_teacher = find_closest_key(query, teacher_schedule_dict.keys())
        response = f"Извините, расписание для {query} не найдено. Вы имели в виду группу {closest_group} или преподавателя {closest_teacher}?"

    safe_send_message(message.chat.id, response) 

def ping_telegram():
    try:
        response = requests.get("https://api.telegram.org")
        return response.status_code == 200
    except:
        return False
    
while True:
    try:
        bot.polling(none_stop=True)
    except Exception as e:
        print(f"Network error: {e}")
        time.sleep(15)






